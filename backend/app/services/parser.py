from pathlib import Path

import fitz
from docx import Document
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models import ParsedDocument, UploadedFile
from app.services.content_safety import (
    check_uploaded_text,
    create_safety_log,
    safety_categories_to_json,
)
from app.services.ids import next_prefixed_id
from app.services.ocr_provider import extract_image_text


IMAGE_TYPES = {"png", "jpg", "jpeg"}


def get_latest_parsed_document(db: Session, file_id: str) -> ParsedDocument | None:
    return db.scalar(
        select(ParsedDocument)
        .where(ParsedDocument.file_id == file_id)
        .order_by(ParsedDocument.created_at.desc())
    )


def ensure_parsed_document(db: Session, uploaded_file: UploadedFile) -> ParsedDocument:
    parsed_document = get_latest_parsed_document(db, uploaded_file.file_id)
    if parsed_document:
        if parsed_document.user_id is None and uploaded_file.user_id:
            parsed_document.user_id = uploaded_file.user_id
            db.flush()
        return parsed_document
    return parse_uploaded_file(db, uploaded_file)


def parse_uploaded_file(db: Session, uploaded_file: UploadedFile) -> ParsedDocument:
    source_name = uploaded_file.storage_key or Path(uploaded_file.file_path).name
    source_path = settings.upload_dir / Path(source_name).name
    parsed_document = ParsedDocument(
        parsed_id=next_prefixed_id(db, ParsedDocument, "parsed"),
        user_id=uploaded_file.user_id,
        project_id=uploaded_file.project_id,
        file_id=uploaded_file.file_id,
        filename=uploaded_file.filename,
        file_type=uploaded_file.file_type,
        parse_status="pending",
    )

    try:
        result = extract_text(source_path, uploaded_file.file_type)
        parsed_document.parse_status = str(result["status"])
        parsed_document.raw_text = str(result["text"])
        parsed_document.text_preview = build_preview(parsed_document.raw_text)
        parsed_document.word_count = len(parsed_document.raw_text)
        parsed_document.ocr_provider = str(result.get("ocr_provider") or "")
        parsed_document.ocr_confidence = float(result.get("ocr_confidence") or 0.0)
        parsed_document.ocr_error = result.get("ocr_error")
        safety_result = check_uploaded_text(
            parsed_document.raw_text,
            {
                "user_id": uploaded_file.user_id,
                "project_id": uploaded_file.project_id,
            },
        )
        parsed_document.safety_warning = safety_result.reason if safety_result.matched_categories else ""
        parsed_document.safety_categories = safety_categories_to_json(
            safety_result.matched_categories
        )
    except Exception as exc:
        parsed_document.parse_status = "failed"
        parsed_document.raw_text = ""
        parsed_document.text_preview = ""
        parsed_document.word_count = 0
        parsed_document.ocr_provider = ""
        parsed_document.ocr_confidence = 0.0
        parsed_document.ocr_error = f"Parse failed: {exc.__class__.__name__}"
        parsed_document.safety_warning = ""
        parsed_document.safety_categories = "[]"

    db.add(parsed_document)
    if parsed_document.safety_warning:
        create_safety_log(
            db,
            result=safety_result,
            direction="uploaded_text",
            input_text=parsed_document.raw_text[:1000],
            user_id=uploaded_file.user_id,
            project_id=uploaded_file.project_id,
        )
    db.commit()
    db.refresh(parsed_document)
    return parsed_document


def extract_text(path: Path, file_type: str) -> dict[str, str | float | None]:
    if file_type == "txt":
        return build_parse_result(read_txt(path), "success")
    if file_type == "docx":
        return build_parse_result(read_docx(path), "success")
    if file_type == "pdf":
        return build_parse_result(read_pdf(path), "success")
    if file_type in IMAGE_TYPES:
        ocr_result = extract_image_text(path)
        return {
            "text": str(ocr_result["text"]),
            "status": str(ocr_result["status"]),
            "ocr_provider": str(ocr_result["provider"]),
            "ocr_confidence": float(ocr_result["confidence"] or 0.0),
            "ocr_error": ocr_result.get("error"),
        }
    raise ValueError(f"Unsupported file type: {file_type}")


def build_parse_result(text: str, status: str) -> dict[str, str | float | None]:
    return {
        "text": text,
        "status": status,
        "ocr_provider": "",
        "ocr_confidence": 0.0,
        "ocr_error": None,
    }


def read_txt(path: Path) -> str:
    raw_bytes = path.read_bytes()
    for encoding in ("utf-8", "gbk"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="replace")


def read_docx(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)


def read_pdf(path: Path) -> str:
    pages: list[str] = []
    with fitz.open(path) as document:
        for index, page in enumerate(document, start=1):
            page_text = page.get_text().strip()
            pages.append(f"[第{index}页]\n{page_text}")
    return "\n\n".join(pages)


def build_preview(raw_text: str, limit: int = 300) -> str:
    compact_text = " ".join(raw_text.split())
    return compact_text[:limit]
